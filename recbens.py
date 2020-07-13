from keras.models import Model, Input
from keras.layers import LSTM, Embedding, Dense, TimeDistributed, Dropout, Bidirectional, Flatten
from keras.callbacks import ModelCheckpoint
from keras.callbacks import EarlyStopping
from keras.models import Sequential
import json
from keras.preprocessing.sequence import pad_sequences
from keras.preprocessing.text import Tokenizer
from keras.preprocessing import sequence
from keras.utils import to_categorical
import docx
import numpy

class Model:

	def __init__(self):
		self.model_accuracy = 96.75
		self.max_len = 30
		self.model_name = 'recbens-model'
		self.num_classes = 3
		self.batch_size = 32
		self.validation_split=0.35
		self.dropout = 0.3
		self.recurrent_dropout = 0.15


	def create_model(self,vocabulary_size,num_classes, max_length,dropout=self.dropout, units=100,recurrent_dropout=self.recurrent_dropout, dense_neurons=16, embedding_vector_length=300):
		inp = Input(shape=(max_length,))
		model = Embedding(input_dim= vocabulary_size+1, output_dim=embedding_vector_length, input_length=max_length)(inp)
		model = Dropout(dropout)(model)
		model = Bidirectional(LSTM(units=units, return_sequences=True, recurrent_dropout=recurrent_dropout))(model)
		#model = TimeDistributed(Dense(dense_neurons, activation='relu'))(model)
		#model = Dense(num_classes, activation="softmax")(model)
		out = TimeDistributed(Dense(3, activation="softmax"))(model)
		#crf = CRF(3, name="output")
		#out = crf(model)
		# softmax output layer
		model = Model(inp,out)
		#model.compile(optimizer="adam", loss=crf.loss_function, metrics=[crf.accuracy],sample_weight_mode="temporal")
		model.compile(optimizer="adam", loss="categorical_crossentropy", metrics=['accuracy'],sample_weight_mode="temporal")
		return model

	# vocabulary_size == len(word2idx)
	def train(self,X_train, X_test, y_train, y_test,c_weights, vocabulary_size, epochs=100,batch_size=self.batch_size,validation_split=self.validation_split, num_classes=self.num_classes, max_length=self.max_len, model_name=self.recbens):
		model = create_model(vocabulary_size,num_classes, max_length)
		print(model.summary())

		es = EarlyStopping(monitor='val_loss', mode='min', verbose=1, patience=10)
		checkpoint = ModelCheckpoint(model_name + '.h5', save_best_only=True, monitor='val_acc', mode='max')

		model.fit(X_train, y_train, epochs=epochs, batch_size=batch_size, validation_split=validation_split,callbacks=[checkpoint,es],verbose=1,sample_weight=c_weights)

		scores = model.evaluate(X_test, y_test, verbose=0)
		print("Accuracy: %.2f%%" % (scores[1] * 100))

		return model


	def load(self,model=None):
		# Loading dictionary
		with open('word_index.json', 'r', encoding='utf-8') as f:
			word2idx = json.load(f)

		# The variable model should include all the path to the model
		if model is None:
			keras_model = create_model(len(word2idx),3,self.max_len)
			keras_model.load_weights('recbens-model-v0-1.h5')
		else:
			keras_model = create_model(len(word2idx),3,self.max_len)
			keras_model.load_weights(model)

		return keras_model

	def load_dictionary(self,dictionary='word_index.json'):
		with open('word_index.json', 'r', encoding='utf-8') as f:
			word2idx = json.load(f)
		return word2idx

	def evaluate_to_docx(self,model,x_sample,text_in_sequences,docx_name='resultado-recbens.docx'):
		doc = docx.Document() # Criando documento docx
		doc_para = doc.add_paragraph(' ') 
		for i in range(len(text_in_sequences)):
			p = []
			p = model.predict(np.array([x_sample[i]]))
			p = np.argmax(p,axis=-1)
			for j in range(len(text_in_sequences[i])):
				doc_para.add_run('  ')
				if p[0][j] == 0:
					doc_para.add_run(text_in_sequences[i][j])
				elif p[0][j] == 1:
					font = doc_para.add_run(text_in_sequences[i][j]).font
					font.highlight_color = 4
				else:
					font = doc_para.add_run(text_in_sequences[i][j]).font
					font.highlight_color = 6
		doc.save(docx_name)

	# From a raw txt, with no preprocessing, this function returns lists with text separed in sequences.
	# One list with the text in sequences with the words changed to respective numbers according to the dictionary and the other list with words not changed
	def preprocess_raw(self,text,word2idx,max_length=self.max_len):
		splitted_text = text.split(' ')

		count = 0
		text_in_sequences = []
		sent_text = []
		text_size = len(splitted_text)
		for i in splitted_text:
			sent_text.append(i)
			count+=1
			if count == max_length:
				text_in_sequences.append(sent_text)
				sent_text = []
				text_size -=max_length
				count = 0
			if(count == text_size):
				text_in_sequences.append(sent_text)

		X_sample = []
		for s in text_in_sequences:
			aux = []
			for w in s:
				if w in word2idx:
					aux.append(word2idx[w])
				else:
					aux.append(word2idx['UNK'])
			X_sample.append(aux) 
		
		x_sample = pad_sequences(sequences = X_sample, maxlen=max_length,value=word2idx["PAD"], padding='post')

		return x_sample,text_in_sequences

	def predict_sequence(self,model,sent_tokens,text_in_sequences,pos,idx2tag):
		p = model.predict(np.array([sent_tokens]))
		p = np.argmax(p,axis=-1)
		sequence_evaluated = []
		for i in range(len(sentences_text[pos])):
			ev = (text_in_sequences[pos][i],idx2tag[p[0][i]])
			sequence_evaluated.append(ev)


	# Given a text, the word dictionary and the class reverse dictionary, this functions returns a list with the classification
	def evaluate(self,model,text,dictionary,class_dictionary = None,docx=False):
		x_sample, text_in_sequences = preprocess_raw(text,dictionary)

		if class_dictionary is None:
			class_dictionary[0] = 'O'
			class_dictionary[1] = 'I-BEM'
			class_dictionary[2] = 'B-BEM'

		if docx == True:
			evaluate_to_docx(model,x_sample,text_in_sequences)
		else:
			full_evaluation = []
			for i in range(len(text_in_sequences)):
				full_evaluation= full_evaluation + predict_sentence(model,x_sample[i],i,class_dictionary)

			return full_evaluation
